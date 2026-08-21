#!/usr/bin/env python3
"""Structure-aware PDF -> DOCX conversion.

pdf2docx reproduces the *appearance* of a PDF using absolute positioning and
direct character formatting. The result looks right and edits terribly: no
Heading styles, no real lists, no navigation pane, no working table of
contents, and reflow that breaks the moment you type.

This converter instead infers document semantics from typography and emits real
Word constructs: Title/Heading 1-3 styles, List Bullet / List Number with true
numbering, real tables, inline images, and character formatting inside runs.

Usage:
    python3 pdf2word-structured.py <input.pdf> <output.docx>

Exit codes:
    0  success
    1  conversion failed
    2  bad arguments
    3  no extractable text layer (caller should OCR first, then retry)
"""

import io
import json
import re
import sys
from collections import Counter, defaultdict

import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# CMR/OT1 fonts leak \textbullet as U+0088 (no ToUnicode map); Word-exported
# PDFs leak Wingdings bullets as private-use U+F0B7 etc. Both are included
# alongside the real Unicode bullets. Whitespace after the marker is optional
# (beamer itemize renders '\u25b6Infer' with no space).
_BULLET_CHARS = (
    "\uf0b7\uf0a7"                 # Wingdings bullet / square
    "\u2022\u2023\u2043\u25aa\u25e6"  # Unicode bullets
    "\u25cf\u25cb\u25a0\u25a1"      # circle / square
    "\u25b6\u25b8\u25ba\u25c0\u00bb"  # triangles / arrows
    "\u2610\u2611\u2713\u2714"      # checkbox / checkmark
    "\u2013\u2014\u00b7"             # en / em dash, middot
    "\x88"                          # OT1 \textbullet leak
)

BULLET_RE = re.compile(r"^\s*([" + _BULLET_CHARS + r"\-\*])\s*(.*)$")
NUMBER_RE = re.compile(r"^\s*(\d{1,3})[.)]\s+(.*)$")
ALPHA_RE = re.compile(r"^\s*([a-zA-Z])[.)]\s+(.*)$")

# A line ending this far short of its block's right edge is a real paragraph
# break rather than wrapped text.
WRAP_TOLERANCE = 0.12
MIN_TEXT_CHARS = 20
TITLE_RATIO = 1.5

# Hebrew + Arabic + RTL punctuation blocks (conservative).
RTL_RE = re.compile(r"[\u0590-\u08ff\ufb1d-\ufdff\ufe70-\ufefc]")

# A line that is only a decorative symbol (e.g. the bullet between an author
# name and date) is dropped rather than becoming its own paragraph.
STANDALONE_SYMBOL_RE = re.compile(r"^[" + _BULLET_CHARS + r"\-\*]+$")

# LaTeX renders section numbers ('1.1', '2') as a separate box from the title,
# so pymupdf reports them as two lines. The number line merges into the title.
SECTION_NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s*$")


def is_bold(span):
    flags = span.get("flags", 0)
    return bool(flags & 2 ** 4) or "bold" in span["font"].lower()


def is_italic(span):
    flags = span.get("flags", 0)
    return (
        bool(flags & 2 ** 1)
        or "italic" in span["font"].lower()
        or "oblique" in span["font"].lower()
    )


def is_rtl_text(text):
    """True when the text is predominantly right-to-left (Hebrew/Arabic)."""
    if not text.strip():
        return False
    rtl = len(RTL_RE.findall(text))
    return rtl / max(len(text), 1) > 0.4


def body_size(doc):
    """Modal font size weighted by character count = the body text size."""
    counter = Counter()
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    counter[round(span["size"], 1)] += len(span["text"].strip())
    return counter.most_common(1)[0][0] if counter else 11.0


def looks_like_heading(text, base, lead):
    """A short bold line at (or slightly above) body size is a heading.

    Size alone misses LaTeX ``subsubsection`` (bold at normalsize) and 12pt
    sections against 11pt body. Requiring a short line that does not end in
    sentence punctuation keeps bold lead-ins inside paragraphs from being
    promoted.
    """
    size = round(lead["size"], 1)
    stripped = text.strip()
    if size > base * 1.12:
        return True
    if size > base * 1.03 and is_bold(lead) and len(stripped) <= 80:
        return not stripped.endswith((".", ",", ";", ":"))
    return False


def heading_scale(doc, base):
    """Map font sizes above body size onto Title / Heading 1..3.

    The largest size that is clearly larger than body text becomes the Title;
    remaining larger sizes map to Heading 1, 2, 3 in descending order (a
    document with a 28pt title and 16pt/14pt headings yields Title, Heading 1,
    Heading 2 — not the Heading-2/Heading-3 skip of the first prototype).
    """
    sizes = set()
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    size = round(span["size"], 1)
                    if size > base * 1.12 and span["text"].strip():
                        sizes.add(size)
    ranked = sorted(sizes, reverse=True)
    levels = {}
    remaining = []
    for index, size in enumerate(ranked):
        if index == 0 and size >= base * TITLE_RATIO:
            levels[size] = "Title"
        else:
            remaining.append(size)
    for index, size in enumerate(remaining):
        levels[size] = f"Heading {min(index + 1, 3)}"
    return levels


def repeated_margins(doc, base):
    """Text at the same vertical position on most pages = header/footer.

    Guarded by font size: body-sized text in a margin is boilerplate, but
    heading-sized text sitting in the top margin is a title, not a header.
    Without that guard, a running header identical to the title deletes the
    title (the first-prototype bug).
    """
    seen = defaultdict(set)
    for number, page in enumerate(doc):
        height = page.rect.height
        for block in page.get_text("dict")["blocks"]:
            if "lines" not in block:
                continue
            top = block["bbox"][1]
            if top > height * 0.08 and top < height * 0.92:
                continue
            spans = [s for l in block["lines"] for s in l["spans"] if s["text"].strip()]
            if not spans:
                continue
            if max(round(s["size"], 1) for s in spans) > base * 1.12:
                continue  # heading-sized text in a margin is a title, not a header
            text = " ".join(s["text"] for s in spans).strip()
            if text:
                seen[re.sub(r"\d+", "#", text)].add(number)
    threshold = max(2, len(doc) * 0.5)
    return {key for key, pages in seen.items() if len(pages) >= threshold}


def is_section_number(text):
    return bool(SECTION_NUMBER_RE.match(text))


def line_text(line):
    return "".join(span["text"] for span in line["spans"])


def dominant(line):
    """The visually dominant span of a line drives its style classification."""
    spans = [s for s in line["spans"] if s["text"].strip()]
    if not spans:
        return None
    return max(spans, key=lambda s: (round(s["size"], 1), len(s["text"])))


def add_runs(paragraph, spans):
    """Preserve bold/italic/colour/rtl as character formatting inside runs."""
    for span in spans:
        text = span["text"]
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = is_bold(span)
        run.italic = is_italic(span)
        colour = span.get("color", 0)
        if colour not in (0, None):
            run.font.color.rgb = RGBColor(
                (colour >> 16) & 255, (colour >> 8) & 255, colour & 255
            )
        if is_rtl_text(text):
            run.font.rtl = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def strip_leading_marker(spans, marker):
    """Remove a list marker (e.g. '1. ' or '\u2022 ') from the first span(s)."""
    out = []
    remaining = marker
    for span in spans:
        text = span["text"]
        if remaining:
            if text.startswith(remaining):
                text = text[len(remaining):]
                remaining = ""
            elif remaining.startswith(text):
                # This span is entirely inside the marker; drop it.
                remaining = remaining[len(text):]
                text = ""
        out.append({**span, "text": text})
    # Drop leading empty/whitespace-only spans so the item text starts clean.
    while out and not out[0]["text"].strip():
        out.pop(0)
    return out


def emit_table(document, table):
    rows = table.extract()
    if not rows or not rows[0]:
        return
    width = max(len(r) for r in rows)
    word_table = document.add_table(rows=len(rows), cols=width)
    word_table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(width):
            value = row[c] if c < len(row) else ""
            cell = word_table.cell(r, c)
            cell.text = (value or "").strip()
            if r == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    document.add_paragraph()


def line_is_wrapped(right, block_right):
    return right >= block_right * (1 - WRAP_TOLERANCE)


def convert(input_path, output_path):
    doc = pymupdf.open(input_path)

    total_text = sum(len(page.get_text().strip()) for page in doc)
    if total_text < MIN_TEXT_CHARS:
        print(
            "No extractable text layer; OCR the PDF first (ocrmypdf) and retry.",
            file=sys.stderr,
        )
        return 3

    base = body_size(doc)
    levels = heading_scale(doc, base)
    boilerplate = repeated_margins(doc, base)

    document = Document()
    document.styles["Normal"].font.size = Pt(base)

    for page_number, page in enumerate(doc):
        if page_number:
            document.add_page_break()

        # Tables first, so their regions can be excluded from the text flow.
        table_regions = []
        try:
            for table in page.find_tables():
                table_regions.append(pymupdf.Rect(table.bbox))
                emit_table(document, table)
        except Exception:
            pass

        pending = None  # (style, spans, left, right)

        def flush():
            nonlocal pending
            if not pending:
                return
            style, spans, _left, _right = pending
            text = "".join(s["text"] for s in spans).strip()
            if text:
                paragraph = document.add_paragraph(style=style)
                if style in ("List Bullet", "List Number"):
                    # Attach live numbering at the paragraph level so the list
                    # renumbers and a w:numPr is visible in document.xml.
                    p_pr = paragraph._p.get_or_add_pPr()
                    num_pr = p_pr.get_or_add_numPr()
                    num_pr.get_or_add_ilvl().val = 0
                    num_pr.get_or_add_numId().val = 1 if style == "List Bullet" else 5
                add_runs(paragraph, spans)
            pending = None

        for block in page.get_text("dict")["blocks"]:
            if "lines" not in block:
                continue
            block_rect = pymupdf.Rect(block["bbox"])
            if any(block_rect.intersects(r) for r in table_regions):
                continue
            block_right = block["bbox"][2]

            for line in block["lines"]:
                text = line_text(line)
                if not text.strip():
                    continue
                if STANDALONE_SYMBOL_RE.match(text.strip()):
                    continue  # decorative separator, not a paragraph
                raw_norm = re.sub(r"\d+", "#", text.strip())
                if raw_norm in boilerplate:
                    continue
                lead = dominant(line)
                if lead is None:
                    continue
                size = round(lead["size"], 1)
                left = line["bbox"][0]
                right = line["bbox"][2]

                style = levels.get(size)
                spans = list(line["spans"])

                if style is None:
                    bullet = BULLET_RE.match(text)
                    number = NUMBER_RE.match(text) or ALPHA_RE.match(text)
                    if bullet:
                        style = "List Bullet"
                        spans = strip_leading_marker(spans, bullet.group(0)[:bullet.start(2)])
                    elif number:
                        style = "List Number"
                        spans = strip_leading_marker(spans, number.group(0)[:number.start(2)])
                    elif looks_like_heading(text, base, lead):
                        style = "Heading 3"
                    else:
                        style = "Normal"

                if style in ("List Bullet", "List Number"):
                    # Each item is its own paragraph; its wrapped continuation
                    # (no marker, indented to the item's left edge) merges.
                    flush()
                    pending = (style, spans, left, right)
                    continue

                if pending and pending[0] in ("List Bullet", "List Number"):
                    # Continuation of a list item.
                    if size <= base * 1.12 and left >= pending[2] - 2:
                        merged = pending[1] + [{**spans[0], "text": " " + spans[0]["text"]}] + spans[1:]
                        pending = (pending[0], merged, pending[2], right)
                        continue
                    flush()

                if style in ("Title",) or style.startswith("Heading"):
                    # Merge a *wrapped* heading, or a section-number line into
                    # the title that follows it (distinct same-style headings
                    # do not merge).
                    if pending and pending[0] == style:
                        pending_text = "".join(s["text"] for s in pending[1]).strip()
                        if line_is_wrapped(pending[3], block_right) or is_section_number(pending_text):
                            merged = pending[1] + [{**spans[0], "text": " " + spans[0]["text"]}] + spans[1:]
                            pending = (style, merged, pending[2], right)
                        else:
                            flush()
                            pending = (style, spans, left, right)
                    else:
                        flush()
                        pending = (style, spans, left, right)
                    continue

                # Normal body text.
                if pending and pending[0] == "Normal" and line_is_wrapped(pending[3], block_right):
                    merged = pending[1] + [{**spans[0], "text": " " + spans[0]["text"]}] + spans[1:]
                    pending = ("Normal", merged, pending[2], right)
                else:
                    flush()
                    pending = ("Normal", spans, left, right)

            flush()

        # Images, placed after the page's text.
        for info in page.get_images(full=True):
            try:
                if info[2] < 20 or info[3] < 20:
                    continue  # skip tiny/background art
                pix = pymupdf.Pixmap(doc, info[0])
                if pix.n - pix.alpha >= 4:
                    pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
                stream = io.BytesIO(pix.tobytes("png"))
                document.add_picture(stream, width=Inches(5.5))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                continue

    document.save(output_path)
    structured_styles = {
        p.style.style_id
        for p in document.paragraphs
        if p.style is not None and p.style.style_id not in ("Normal", "No List")
    }
    text_chars = sum(len(p.text) for p in document.paragraphs)
    print(
        "STRUCTURED_STATS "
        + json.dumps({"styles": len(structured_styles), "text_chars": text_chars}),
        flush=True,
    )
    return 0


def main():
    if len(sys.argv) != 3:
        print("usage: pdf2word-structured.py <input.pdf> <output.docx>", file=sys.stderr)
        return 2
    try:
        return convert(sys.argv[1], sys.argv[2])
    except Exception as exc:  # noqa: BLE001
        print(f"structured conversion failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
